import json
import re
import os
from datetime import datetime
from collections import defaultdict
try:
    import boto3
except ImportError:
    boto3 = None
try:
    from docx import Document
except ImportError:
    print("Warning: python-docx not installed. Document processing may fail.")
    Document = None

class DocumentAnalyzer:
    def __init__(self):
        self.hawkeye_sections = {
            1: "Initial Assessment",
            2: "Investigation Process", 
            3: "Seller Classification",
            4: "Enforcement Decision-Making",
            5: "Additional Verification (High-Risk Cases)",
            6: "Multiple Appeals Handling",
            7: "Account Hijacking Prevention",
            8: "Funds Management",
            9: "REs-Q Outreach Process",
            10: "Sentiment Analysis",
            11: "Root Cause Analysis",
            12: "Preventative Actions",
            13: "Documentation and Reporting",
            14: "Cross-Team Collaboration",
            15: "Quality Control",
            16: "Continuous Improvement",
            17: "Communication Standards",
            18: "Performance Metrics",
            19: "Legal and Compliance",
            20: "New Service Launch Considerations"
        }
        
        self.standard_sections = [
            "Executive Summary", "Background", "Timeline of Events",
            "Resolving Actions", "Root Causes (RC) and Preventative Actions (PA)",
            "Root Cause", "Preventative Actions", "Investigation Process",
            "Seller Classification", "Documentation and Reporting",
            "Impact Assessment", "Timeline", "Recommendations"
        ]
        
        self.excluded_sections = [
            "Original Email", "Email Correspondence", "Raw Data", "Logs",
            "Attachments", "From:", "Sent:", "To:", "Cc:", "Subject:"
        ]

    def extract_sections_from_docx(self, doc_path):
        """Extract sections from Word document with comprehensive content capture"""
        try:
            if Document is None:
                raise ImportError("python-docx not available")
            
            if not os.path.exists(doc_path):
                raise FileNotFoundError(f"Document not found: {doc_path}")
            
            print(f"Loading document: {doc_path}")
            doc = Document(doc_path)
            print(f"Document loaded successfully")
            
            sections = {}
            section_paragraphs = {}
            paragraph_indices = {}
            
            # First try header-based detection
            sections, section_paragraphs, paragraph_indices = self._extract_by_headers(doc)
            
            # If insufficient sections found, try AI-based detection
            if len(sections) < 3:
                print(f"Only {len(sections)} sections found, trying AI detection...")
                ai_sections = self._identify_sections_with_ai(doc)
                if ai_sections:
                    sections, section_paragraphs, paragraph_indices = self._extract_by_ai_hints(doc, ai_sections)
            
            # Fallback: create single section with all content
            if not sections:
                print(f"No sections detected, creating single section...")
                sections, section_paragraphs, paragraph_indices = self._create_single_section(doc)
            
            print(f"Extracted {len(sections)} sections: {list(sections.keys())}")
            return sections, section_paragraphs, paragraph_indices
            
        except Exception as e:
            print(f"Document loading failed: {str(e)}")
            # Return safe fallback structure
            return {
                "Document": f"Failed to load document: {str(e)}"
            }, {
                "Document": [f"Failed to load document: {str(e)}"]
            }, {
                "Document": [0]
            }

    def _extract_by_headers(self, doc):
        """Extract sections using header detection"""
        sections = {}
        section_paragraphs = {}
        paragraph_indices = {}
        
        all_paragraphs = [(idx, para, para.text.strip()) for idx, para in enumerate(doc.paragraphs) if para.text.strip()]
        section_headers = []
        
        # Find section headers
        for idx, para, text in all_paragraphs:
            if len(text) < 100:  # Headers are typically short
                for section_name in self.standard_sections:
                    if section_name.lower() in text.lower():
                        section_headers.append({'title': section_name, 'idx': idx})
                        break
                
                # Pattern-based header detection
                if re.match(r'^(\d+\.? )?[A-Z][a-z]+( [A-Z][a-z]+){0,3}$', text) and len(text.split()) <= 5:
                    section_headers.append({'title': text, 'idx': idx})
        
        # Extract content between headers
        section_headers.sort(key=lambda x: x['idx'])
        
        for i, header in enumerate(section_headers):
            section_title = header['title']
            start_idx = header['idx']
            end_idx = len(doc.paragraphs)
            
            if i < len(section_headers) - 1:
                end_idx = section_headers[i+1]['idx']
            
            content, paras, indices = self._extract_section_content(doc, start_idx, end_idx)
            
            if content:
                sections[section_title] = content
                section_paragraphs[section_title] = paras
                paragraph_indices[section_title] = indices
        
        return sections, section_paragraphs, paragraph_indices

    def _identify_sections_with_ai(self, doc):
        """Use AI to identify document sections"""
        full_text = '\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        
        prompt = f"""You are a senior document structure analyst with comprehensive expertise in business document organization, professional investigation reports, and CT EE investigation documentation standards. Your specialized task is to systematically identify and extract all main sections from this professional investigation document using established analytical frameworks.

DOCUMENT TEXT FOR SYSTEMATIC ANALYSIS (first 10,000 characters):
{full_text[:10000]}

COMPREHENSIVE SECTION IDENTIFICATION METHODOLOGY:
1. Content Transition Analysis - Identify clear content transitions and investigative topic changes that indicate distinct section boundaries
2. Header Recognition - Locate headers that appear on dedicated lines or introduce new investigative topics, methodologies, or findings
3. Standard Section Detection - Find common professional document sections (Executive Summary, Timeline, Background, Analysis, etc.)
4. Functional Heading Identification - Detect text that functions as organizational headings even without explicit formatting
5. Structural Pattern Recognition - Look for numbered, bulleted, or hierarchical section starts that organize investigative content
6. Chronological Content Mapping - Identify date-based sections or chronological content blocks that structure the investigation
7. Investigation-Specific Section Recognition - Recognize specialized investigation sections (Root Cause, Preventative Actions, Impact Assessment, etc.)

PROFESSIONAL CT EE INVESTIGATION SECTION PATTERNS (COMPREHENSIVE LIST):
- Executive Summary / Investigation Summary / Key Findings Summary
- Background / Context / Issue Description / Problem Statement
- Timeline of Events / Chronology / Event Sequence / Incident Timeline
- Investigation Process / Methodology / Analysis Approach / Investigation Framework
- Findings / Results / Conclusions / Investigation Outcomes
- Resolving Actions / Remediation Steps / Actions Taken / Corrective Measures
- Root Causes (RC) and Preventative Actions (PA) / Root Cause Analysis
- Impact Assessment / Business Impact / Customer Impact / Operational Impact
- Recommendations / Next Steps / Future Actions / Strategic Recommendations
- Conclusion / Closing / Lessons Learned / Investigation Closure
- Stakeholder Communication / Notification Process / Escalation Procedures
- Quality Assurance / Validation / Verification Process

SYSTEMATIC ANALYSIS INSTRUCTIONS:
1. Document Scanning - Systematically scan the document text from beginning to end using professional document analysis techniques
2. Boundary Identification - Identify clear section boundaries where investigative topics, methodologies, or focus areas change
3. Title Extraction - Extract the exact section title from the document when clearly identifiable, maintaining professional terminology
4. Distinctive Phrase Identification - Find a unique, distinctive phrase from the beginning of each section as a "line_hint" for precise location
5. Sequential Organization - Ensure sections are listed in the exact order they appear in the original document structure
6. Content Validation - Focus exclusively on substantial investigative content sections (exclude headers, footers, metadata, administrative content)

STRICT JSON OUTPUT SPECIFICATION:
{{
    "sections": [
        {{"title": "Professional Section Name", "line_hint": "distinctive opening phrase from section content"}},
        {{"title": "Next Section Title", "line_hint": "unique identifying phrase from section start"}}
    ]
}}

MANDATORY COMPLIANCE REQUIREMENTS:
- Return EXCLUSIVELY valid JSON format with no additional text, explanations, or commentary before or after
- Use exact section titles from the document when clearly identifiable (maintain professional terminology and formatting)
- Line hints must be unique, distinctive phrases that clearly and unambiguously identify section starting points
- Include ALL meaningful investigative sections found (minimum 2 sections, maximum 12 sections for comprehensive coverage)
- Maintain the original document sequential order exactly as it appears in the source material
- Ensure each identified section contains substantial, meaningful investigative content (exclude administrative or formatting elements)
- Prioritize professional investigation terminology and standard section naming conventions"""
        
        try:
            response = self._invoke_bedrock("You are an expert document structure analyst with extensive experience in business document organization and content identification. You excel at recognizing section boundaries, content transitions, and organizational patterns in professional documents, even when sections lack explicit formatting or clear headers.", prompt)
            result = json.loads(response)
            return result.get('sections', [])
        except:
            return None

    def _extract_by_ai_hints(self, doc, ai_sections):
        """Extract sections using AI-identified hints"""
        sections = {}
        section_paragraphs = {}
        paragraph_indices = {}
        
        all_sections_info = []
        
        for section_info in ai_sections:
            section_title = section_info.get('title', '')
            line_hint = section_info.get('line_hint', '').lower()
            
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip().lower()
                
                if line_hint and line_hint in text:
                    all_sections_info.append({'title': section_title, 'start_idx': idx})
                    break
                elif section_title.lower() in text:
                    all_sections_info.append({'title': section_title, 'start_idx': idx})
                    break
        
        all_sections_info.sort(key=lambda x: x['start_idx'])
        
        for i, section_info in enumerate(all_sections_info):
            section_title = section_info['title']
            start_idx = section_info['start_idx']
            end_idx = len(doc.paragraphs)
            
            if i < len(all_sections_info) - 1:
                end_idx = all_sections_info[i+1]['start_idx']
            
            content, paras, indices = self._extract_section_content(doc, start_idx, end_idx)
            
            if content:
                sections[section_title] = content
                section_paragraphs[section_title] = paras
                paragraph_indices[section_title] = indices
        
        return sections, section_paragraphs, paragraph_indices

    def _create_single_section(self, doc):
        """Create single section with all content as fallback"""
        sections = {}
        section_paragraphs = {}
        paragraph_indices = {}
        
        content = []
        paras = []
        indices = []
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                content.append(text)
                paras.append(para)
                indices.append(idx)
        
        if content:
            sections["Document Content"] = '\n\n'.join(content)
            section_paragraphs["Document Content"] = paras
            paragraph_indices["Document Content"] = indices
        
        return sections, section_paragraphs, paragraph_indices

    def _extract_section_content(self, doc, start_idx, end_idx):
        """Extract content between two paragraph indices"""
        content = []
        paras = []
        indices = []
        
        for idx in range(start_idx, end_idx):
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                text = para.text.strip()
                
                if idx == start_idx or not text:
                    continue
                
                # Skip email dividers
                if any(text.startswith(prefix) for prefix in ["From:", "Sent:", "To:", "---"]):
                    continue
                
                content.append(text)
                paras.append(para)
                indices.append(idx)
        
        return '\n\n'.join(content), paras, indices

    def _invoke_bedrock(self, system_prompt, user_prompt):
        """Invoke AWS Bedrock for AI analysis"""
        try:
            if boto3 is None:
                raise ImportError("boto3 not available")
                
            runtime = boto3.client('bedrock-runtime')
            
            body = json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 4000,
                "system": system_prompt,
                "messages": [{"role": "user", "content": user_prompt}]
            })
            
            response = runtime.invoke_model(
                body=body,
                modelId='anthropic.claude-3-sonnet-20240229-v1:0',
                accept="application/json",
                contentType="application/json"
            )
            
            response_body = json.loads(response.get('body').read())
            return response_body['content'][0]['text']
            
        except Exception as e:
            print(f"AI section detection failed: {str(e)}")
            # Fallback response for testing
            return json.dumps({
                "sections": [
                    {"title": "Executive Summary", "line_hint": "executive summary"},
                    {"title": "Timeline of Events", "line_hint": "timeline"},
                    {"title": "Resolving Actions", "line_hint": "resolving actions"},
                    {"title": "Root Causes (RC) and Preventative Actions (PA)", "line_hint": "root cause"}
                ]
            })