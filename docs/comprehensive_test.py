#!/usr/bin/env python3
"""
Comprehensive End-to-End Testing Suite for AI-Prism
Tests all major functionality paths and generates detailed report
"""

import requests
import json
import time
from datetime import datetime
from pathlib import Path
import sys

# Test Configuration
BASE_URL = "http://localhost:8080"
TEST_RESULTS = []
TEST_DOCUMENT = "test_sample.docx"

class TestResult:
    def __init__(self, test_name, category):
        self.test_name = test_name
        self.category = category
        self.status = "PENDING"
        self.duration = 0
        self.details = []
        self.errors = []
        self.start_time = None

    def start(self):
        self.start_time = time.time()
        print(f"\n{'='*80}")
        print(f"🧪 TEST: {self.test_name}")
        print(f"{'='*80}")

    def add_detail(self, detail):
        self.details.append(detail)
        print(f"  ✓ {detail}")

    def add_error(self, error):
        self.errors.append(error)
        print(f"  ✗ ERROR: {error}")

    def complete(self, status):
        self.duration = time.time() - self.start_time
        self.status = status
        status_emoji = "✅" if status == "PASS" else "❌" if status == "FAIL" else "⚠️"
        print(f"\n{status_emoji} {self.status} ({self.duration:.2f}s)")
        TEST_RESULTS.append(self)

def test_server_status():
    """Test 1: Server Connectivity"""
    test = TestResult("Server Status and Connectivity", "Infrastructure")
    test.start()

    try:
        response = requests.get(f"{BASE_URL}/", timeout=5)
        test.add_detail(f"Server response: {response.status_code}")

        if response.status_code == 200:
            test.add_detail("Server is running and accessible")
            test.add_detail(f"Response size: {len(response.text)} bytes")
            test.complete("PASS")
            return True
        else:
            test.add_error(f"Unexpected status code: {response.status_code}")
            test.complete("FAIL")
            return False
    except Exception as e:
        test.add_error(f"Connection failed: {str(e)}")
        test.complete("FAIL")
        return False

def test_claude_connection():
    """Test 2: Claude API Connection"""
    test = TestResult("Claude API Connection Test", "Infrastructure")
    test.start()

    try:
        response = requests.get(f"{BASE_URL}/test_claude_connection", timeout=10)
        data = response.json()

        test.add_detail(f"Response status: {response.status_code}")
        test.add_detail(f"Connected: {data.get('connected', False)}")

        if data.get('connected'):
            test.add_detail(f"Model: {data.get('model', 'Unknown')}")
            test.add_detail(f"Response time: {data.get('response_time', 0):.2f}s")
            test.complete("PASS")
            return True
        else:
            test.add_error(f"Claude connection failed: {data.get('error', 'Unknown error')}")
            test.complete("FAIL")
            return False
    except Exception as e:
        test.add_error(f"Test failed: {str(e)}")
        test.complete("FAIL")
        return False

def create_test_document():
    """Helper: Create a test Word document"""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading('Test Investigation Report', 0)

        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(
            'This is a test investigation report for AI-Prism end-to-end testing. '
            'The case involves a seller account verification issue that required immediate attention.'
        )

        doc.add_heading('Timeline of Events', level=1)
        doc.add_paragraph('• 2024-01-15: Initial complaint received')
        doc.add_paragraph('• 2024-01-16: Investigation initiated')
        doc.add_paragraph('• 2024-01-17: Evidence collected and reviewed')

        doc.add_heading('Root Cause Analysis', level=1)
        doc.add_paragraph(
            'The root cause was identified as a system configuration error that prevented '
            'proper validation of seller credentials during the onboarding process.'
        )

        doc.add_heading('Preventative Actions', level=1)
        doc.add_paragraph('1. Update validation rules in the system')
        doc.add_paragraph('2. Implement additional verification checkpoints')
        doc.add_paragraph('3. Train support team on new procedures')

        doc.save(TEST_DOCUMENT)
        return True
    except Exception as e:
        print(f"Error creating test document: {e}")
        return False

def test_document_upload():
    """Test 3: Document Upload and Section Extraction"""
    test = TestResult("Document Upload and Section Extraction", "Core Functionality")
    test.start()

    # Create test document
    if not create_test_document():
        test.add_error("Failed to create test document")
        test.complete("FAIL")
        return None

    test.add_detail("Test document created successfully")

    try:
        with open(TEST_DOCUMENT, 'rb') as f:
            files = {'document': (TEST_DOCUMENT, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'guidelines_preference': 'both'}

            response = requests.post(f"{BASE_URL}/upload", files=files, data=data, timeout=30)
            result = response.json()

            test.add_detail(f"Upload response: {response.status_code}")

            if result.get('success'):
                session_id = result.get('session_id')
                sections = result.get('sections', [])

                test.add_detail(f"Session ID: {session_id}")
                test.add_detail(f"Sections extracted: {len(sections)}")
                test.add_detail(f"Section names: {', '.join(sections)}")

                if len(sections) >= 3:  # Should extract at least 3 sections
                    test.complete("PASS")
                    return session_id
                else:
                    test.add_error(f"Expected at least 3 sections, got {len(sections)}")
                    test.complete("WARN")
                    return session_id
            else:
                test.add_error(f"Upload failed: {result.get('error', 'Unknown error')}")
                test.complete("FAIL")
                return None
    except Exception as e:
        test.add_error(f"Test failed: {str(e)}")
        test.complete("FAIL")
        return None

def test_section_analysis(session_id):
    """Test 4: On-Demand Section Analysis"""
    test = TestResult("On-Demand Section Analysis", "Core Functionality")
    test.start()

    if not session_id:
        test.add_error("No session ID provided")
        test.complete("FAIL")
        return None

    try:
        # Analyze first section
        payload = {
            'session_id': session_id,
            'section_name': 'Executive Summary'
        }

        response = requests.post(
            f"{BASE_URL}/analyze_section",
            json=payload,
            headers={'Content-Type': 'application/json'},
            timeout=60
        )
        data = response.json()

        test.add_detail(f"Analysis response: {response.status_code}")
        test.add_detail(f"Async mode: {data.get('async', False)}")

        if data.get('async') and data.get('task_id'):
            task_id = data['task_id']
            test.add_detail(f"Task ID: {task_id}")
            test.add_detail(f"Section content included: {bool(data.get('section_content'))}")

            # Poll for completion
            test.add_detail("Polling for task completion...")
            max_attempts = 60
            for attempt in range(max_attempts):
                time.sleep(2)
                status_response = requests.get(f"{BASE_URL}/task_status/{task_id}")
                status_data = status_response.json()

                state = status_data.get('state')
                test.add_detail(f"Poll attempt {attempt + 1}: {state}")

                if state == 'SUCCESS':
                    result = status_data.get('result', {})

                    # ✅ FIXED: Add debug output to diagnose parsing issue
                    test.add_detail(f"Raw result keys: {list(result.keys()) if isinstance(result, dict) else 'not a dict'}")
                    test.add_detail(f"Result success: {result.get('success', 'N/A')}")

                    feedback_items = result.get('feedback_items', [])

                    test.add_detail(f"Analysis complete!")
                    test.add_detail(f"Feedback items type: {type(feedback_items)}")
                    test.add_detail(f"Feedback items generated: {len(feedback_items) if isinstance(feedback_items, list) else 'N/A'}")
                    test.add_detail(f"Duration: {result.get('duration', 0):.2f}s")
                    test.add_detail(f"Model used: {result.get('model_used', 'Unknown')}")

                    if isinstance(feedback_items, list) and len(feedback_items) > 0:
                        test.add_detail(f"First feedback item keys: {list(feedback_items[0].keys())[:5]}")
                        test.complete("PASS")
                        return feedback_items
                    else:
                        if not isinstance(feedback_items, list):
                            test.add_error(f"feedback_items is not a list: {type(feedback_items)}")
                        else:
                            test.add_error("No feedback items generated (empty list)")
                        test.complete("WARN")
                        return []
                elif state == 'FAILURE':
                    test.add_error(f"Task failed: {status_data.get('error', 'Unknown error')}")
                    test.complete("FAIL")
                    return None

            test.add_error("Task polling timeout")
            test.complete("FAIL")
            return None
        else:
            test.add_error("Expected async response with task_id")
            test.complete("FAIL")
            return None
    except Exception as e:
        test.add_error(f"Test failed: {str(e)}")
        test.complete("FAIL")
        return None

def test_chat_functionality(session_id):
    """Test 5: Chat Functionality"""
    test = TestResult("Chat Assistant Functionality", "User Interaction")
    test.start()

    if not session_id:
        test.add_error("No session ID provided")
        test.complete("FAIL")
        return False

    try:
        payload = {
            'message': 'What are the main Hawkeye checkpoints for Executive Summary?',
            'session_id': session_id,
            'current_section': 'Executive Summary'
        }

        response = requests.post(
            f"{BASE_URL}/chat",
            json=payload,
            headers={'Content-Type': 'application/json'},
            timeout=60
        )
        data = response.json()

        test.add_detail(f"Chat response: {response.status_code}")
        test.add_detail(f"Async mode: {data.get('async', False)}")

        if data.get('async') and data.get('task_id'):
            task_id = data['task_id']
            test.add_detail(f"Task ID: {task_id}")

            # Poll for completion
            test.add_detail("Polling for chat response...")
            max_attempts = 30
            for attempt in range(max_attempts):
                time.sleep(2)
                status_response = requests.get(f"{BASE_URL}/task_status/{task_id}")
                status_data = status_response.json()

                state = status_data.get('state')

                if state == 'SUCCESS':
                    result = status_data.get('result', {})
                    response_text = result.get('response', '')

                    test.add_detail(f"Chat response received!")
                    test.add_detail(f"Response length: {len(response_text)} characters")
                    test.add_detail(f"Model used: {result.get('model_used', 'Unknown')}")

                    if len(response_text) > 50:
                        test.add_detail(f"Preview: {response_text[:100]}...")
                        test.complete("PASS")
                        return True
                    else:
                        test.add_error("Response too short or empty")
                        test.complete("WARN")
                        return False
                elif state == 'FAILURE':
                    test.add_error(f"Chat failed: {status_data.get('error', 'Unknown error')}")
                    test.complete("FAIL")
                    return False

            test.add_error("Chat polling timeout")
            test.complete("FAIL")
            return False
        else:
            test.add_error("Expected async response with task_id")
            test.complete("FAIL")
            return False
    except Exception as e:
        test.add_error(f"Test failed: {str(e)}")
        test.complete("FAIL")
        return False

def test_statistics_endpoint(session_id):
    """Test 6: Statistics Endpoint"""
    test = TestResult("Statistics and Session Management", "Data Management")
    test.start()

    if not session_id:
        test.add_error("No session ID provided")
        test.complete("FAIL")
        return False

    try:
        response = requests.get(f"{BASE_URL}/get_statistics?session_id={session_id}", timeout=10)
        data = response.json()

        test.add_detail(f"Statistics response: {response.status_code}")
        test.add_detail(f"Sections analyzed: {data.get('sections_analyzed', 0)}")
        test.add_detail(f"Total feedback: {data.get('total_feedback', 0)}")
        test.add_detail(f"High risk items: {data.get('high_risk_count', 0)}")
        test.add_detail(f"Medium risk items: {data.get('medium_risk_count', 0)}")

        if response.status_code == 200:
            test.complete("PASS")
            return True
        else:
            test.add_error("Statistics endpoint returned unexpected status")
            test.complete("FAIL")
            return False
    except Exception as e:
        test.add_error(f"Test failed: {str(e)}")
        test.complete("FAIL")
        return False

def generate_report():
    """Generate comprehensive test report"""
    print("\n\n" + "="*80)
    print("📊 COMPREHENSIVE TEST REPORT")
    print("="*80)
    print(f"Test Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Base URL: {BASE_URL}")
    print(f"Total Tests: {len(TEST_RESULTS)}")

    # Summary by status
    passed = sum(1 for t in TEST_RESULTS if t.status == "PASS")
    failed = sum(1 for t in TEST_RESULTS if t.status == "FAIL")
    warned = sum(1 for t in TEST_RESULTS if t.status == "WARN")

    print(f"\n✅ PASSED: {passed}")
    print(f"❌ FAILED: {failed}")
    print(f"⚠️  WARNINGS: {warned}")

    # Detailed results by category
    categories = {}
    for test in TEST_RESULTS:
        if test.category not in categories:
            categories[test.category] = []
        categories[test.category].append(test)

    print("\n" + "="*80)
    print("DETAILED RESULTS BY CATEGORY")
    print("="*80)

    for category, tests in categories.items():
        print(f"\n📁 {category}")
        print("-" * 80)
        for test in tests:
            status_emoji = "✅" if test.status == "PASS" else "❌" if test.status == "FAIL" else "⚠️"
            print(f"{status_emoji} {test.test_name} ({test.duration:.2f}s)")

            if test.details:
                for detail in test.details[:5]:  # Show first 5 details
                    print(f"    • {detail}")

            if test.errors:
                for error in test.errors:
                    print(f"    ✗ {error}")

    # Save report to file
    report_file = f"test_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    with open(report_file, 'w') as f:
        f.write("="*80 + "\n")
        f.write("COMPREHENSIVE END-TO-END TEST REPORT\n")
        f.write("="*80 + "\n")
        f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Base URL: {BASE_URL}\n\n")

        for category, tests in categories.items():
            f.write(f"\n{category}\n")
            f.write("-" * 80 + "\n")
            for test in tests:
                f.write(f"\n{test.status}: {test.test_name} ({test.duration:.2f}s)\n")
                if test.details:
                    f.write("  Details:\n")
                    for detail in test.details:
                        f.write(f"    - {detail}\n")
                if test.errors:
                    f.write("  Errors:\n")
                    for error in test.errors:
                        f.write(f"    - {error}\n")

    print(f"\n📄 Full report saved to: {report_file}")

    return passed == len(TEST_RESULTS)

def main():
    """Run all tests"""
    print("\n" + "="*80)
    print("🚀 AI-PRISM COMPREHENSIVE END-TO-END TEST SUITE")
    print("="*80)

    # Test 1: Server Status
    if not test_server_status():
        print("\n❌ Server not accessible. Aborting tests.")
        return False

    # Test 2: Claude Connection
    test_claude_connection()

    # Test 3: Document Upload
    session_id = test_document_upload()

    # Test 4: Section Analysis
    feedback_items = test_section_analysis(session_id) if session_id else None

    # Test 5: Chat Functionality
    test_chat_functionality(session_id) if session_id else None

    # Test 6: Statistics
    test_statistics_endpoint(session_id) if session_id else None

    # Generate Report
    all_passed = generate_report()

    if all_passed:
        print("\n✅ ALL TESTS PASSED!")
        return True
    else:
        print("\n⚠️  SOME TESTS FAILED - See report for details")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
