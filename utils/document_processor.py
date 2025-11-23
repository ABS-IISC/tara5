import os
import json
import zipfile
import shutil
import uuid
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from lxml import etree

class DocumentProcessor:
    def __init__(self):
        self.temp_dirs = []

    def create_document_with_comments(self, original_path, comments_data, output_filename=None):
        """Create a Word document with proper comments"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"reviewed_document_{timestamp}.docx"

        # DEBUG: Log document processing
        print(f"\n{'='*60}")
        print(f"📄 DOCUMENT PROCESSOR - create_document_with_comments")
        print(f"{'='*60}")
        print(f"Original document: {original_path}")
        print(f"Output filename: {output_filename}")
        print(f"Number of comments to insert: {len(comments_data)}")
        print(f"Comments data received: {comments_data is not None}")
        if comments_data:
            print(f"Sample comment: {comments_data[0] if len(comments_data) > 0 else 'None'}")
        print(f"{'='*60}\n")

        try:
            # Use the advanced comment insertion method
            print("🔧 Attempting XML comment insertion method...")
            result = self._create_with_xml_comments(original_path, comments_data, output_filename)
            print(f"✅ XML comment method succeeded: {result}")
            return result
        except Exception as e:
            print(f"❌ XML comment method failed: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to annotation method
            print("🔧 Falling back to annotation method...")
            result = self._create_with_annotations(original_path, comments_data, output_filename)
            print(f"✅ Annotation method result: {result}")
            return result

    def _create_with_xml_comments(self, original_path, comments_data, output_filename):
        """Create document with XML-based comments"""
        temp_dir = f"temp_{uuid.uuid4()}"
        self.temp_dirs.append(temp_dir)

        print(f"📦 XML Method - Creating temp directory: {temp_dir}")

        try:
            # Create a copy of the original document
            temp_docx = f"{temp_dir}_temp.docx"
            print(f"📋 Copying original document to: {temp_docx}")
            doc = Document(original_path)
            doc.save(temp_docx)

            # Extract the docx as a zip
            os.makedirs(temp_dir, exist_ok=True)
            print(f"📂 Extracting docx to: {temp_dir}")
            with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # Create comments.xml
            print(f"💬 Generating comments.xml with {len(comments_data)} comments...")
            comments_xml = self._generate_comments_xml(comments_data)
            comments_path = os.path.join(temp_dir, 'word', 'comments.xml')

            with open(comments_path, 'w', encoding='utf-8') as f:
                f.write(comments_xml)
            print(f"✅ Created comments.xml at: {comments_path}")

            # Update document.xml.rels
            print(f"🔗 Updating document.xml.rels...")
            self._update_document_rels(temp_dir)

            # Update [Content_Types].xml
            print(f"📑 Updating [Content_Types].xml...")
            self._update_content_types(temp_dir)

            # Update document.xml with comment references
            print(f"🔖 Inserting comment references into document.xml...")
            self._insert_comment_references(temp_dir, comments_data)

            # Repackage as docx
            print(f"📦 Repackaging as docx: {output_filename}")
            with zipfile.ZipFile(output_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, temp_dir)
                        zipf.write(file_path, arcname)

            print(f"✅ Successfully created document with {len(comments_data)} comments")
            return output_filename

        finally:
            # Cleanup
            print(f"🧹 Cleaning up temporary files...")
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

    def _generate_comments_xml(self, comments_data):
        """Generate the comments.xml content"""
        xml_content = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'''
        
        for i, comment in enumerate(comments_data):
            comment_id = i + 1
            author = comment.get('author', 'AI Feedback')
            date = datetime.now().strftime('%Y-%m-%dT%H:%M:%S.%fZ')
            text = comment.get('comment', '')
            
            xml_content += f'''
    <w:comment w:id="{comment_id}" w:author="{author}" w:date="{date}">
        <w:p>
            <w:r>
                <w:t>{text}</w:t>
            </w:r>
        </w:p>
    </w:comment>'''
        
        xml_content += '\n</w:comments>'
        return xml_content

    def _update_document_rels(self, temp_dir):
        """Update document.xml.rels to include comments relationship"""
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        
        if os.path.exists(rels_path):
            with open(rels_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if 'comments.xml' not in content:
                # Find the highest rId number
                import re
                rids = re.findall(r'rId(\d+)', content)
                max_rid = max([int(rid) for rid in rids]) if rids else 0
                new_rid = f"rId{max_rid + 1}"
                
                new_rel = f'<Relationship Id="{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                content = content.replace('</Relationships>', f'{new_rel}</Relationships>')
                
                with open(rels_path, 'w', encoding='utf-8') as f:
                    f.write(content)

    def _update_content_types(self, temp_dir):
        """Update [Content_Types].xml to include comments content type"""
        content_types_path = os.path.join(temp_dir, '[Content_Types].xml')
        
        if os.path.exists(content_types_path):
            with open(content_types_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if 'comments.xml' not in content:
                new_type = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                content = content.replace('</Types>', f'{new_type}</Types>')
                
                with open(content_types_path, 'w', encoding='utf-8') as f:
                    f.write(content)

    def _insert_comment_references(self, temp_dir, comments_data):
        """Insert comment references into document.xml"""
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')

        print(f"🔖 Inserting {len(comments_data)} comment references...")
        print(f"   Document path: {doc_xml_path}")

        if not os.path.exists(doc_xml_path):
            print(f"❌ ERROR: document.xml not found at {doc_xml_path}")
            return

        # Parse the document XML
        tree = etree.parse(doc_xml_path)
        root = tree.getroot()

        # Find paragraphs and insert comment references
        paragraphs = root.xpath('//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        print(f"   Found {len(paragraphs)} paragraphs in document")

        for i, comment in enumerate(comments_data):
            comment_id = i + 1
            para_index = comment.get('paragraph_index', 0)

            print(f"\n   Comment {comment_id}:")
            print(f"      Target paragraph: {para_index}")
            print(f"      Section: {comment.get('section', 'Unknown')}")
            print(f"      Comment preview: {comment.get('comment', '')[:50]}...")

            if para_index < len(paragraphs):
                para = paragraphs[para_index]

                # Create comment range start
                comment_start = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
                comment_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))

                # Create comment range end
                comment_end = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                comment_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))

                # Create comment reference
                comment_ref = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                comment_ref_elem = etree.SubElement(comment_ref, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
                comment_ref_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))

                # Insert at the beginning of the paragraph
                para.insert(0, comment_start)
                para.append(comment_end)
                para.append(comment_ref)

                print(f"      ✅ Inserted comment reference at paragraph {para_index}")
            else:
                print(f"      ⚠️ WARNING: Paragraph index {para_index} out of range (max: {len(paragraphs)-1})")

        # Save the modified document.xml
        print(f"\n💾 Saving modified document.xml...")
        tree.write(doc_xml_path, encoding='utf-8', xml_declaration=True)
        print(f"✅ document.xml saved successfully")

    def _create_with_annotations(self, original_path, comments_data, output_filename):
        """Fallback method: create document with inline annotations"""
        try:
            original_doc = Document(original_path)
            new_doc = Document()
            
            # Copy document structure and add annotations
            for para_idx, para in enumerate(original_doc.paragraphs):
                # Copy paragraph
                new_para = new_doc.add_paragraph()
                new_para.style = para.style
                
                # Copy runs
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    if run.bold:
                        new_run.bold = True
                    if run.italic:
                        new_run.italic = True
                    if run.underline:
                        new_run.underline = True
                
                # Add comments for this paragraph
                para_comments = [c for c in comments_data if c.get('paragraph_index') == para_idx]
                
                for comment in para_comments:
                    comment_para = new_doc.add_paragraph()
                    comment_run = comment_para.add_run(f"💬 {comment.get('author', 'AI Feedback')}: {comment.get('comment', '')}")
                    comment_run.font.color.rgb = RGBColor(102, 126, 234)
                    comment_run.font.size = Pt(10)
                    comment_run.italic = True
            
            # Add summary section
            self._add_feedback_summary(new_doc, comments_data)
            
            new_doc.save(output_filename)
            return output_filename
            
        except Exception as e:
            print(f"Annotation method failed: {e}")
            return self._create_simple_copy(original_path, comments_data, output_filename)

    def _create_simple_copy(self, original_path, comments_data, output_filename):
        """Simple fallback: copy document and add summary"""
        try:
            doc = Document(original_path)
            
            # Add feedback summary at the end
            self._add_feedback_summary(doc, comments_data)
            
            doc.save(output_filename)
            return output_filename
            
        except Exception as e:
            print(f"Simple copy method failed: {e}")
            return None

    def _add_feedback_summary(self, doc, comments_data):
        """Add comprehensive feedback summary to document"""
        doc.add_page_break()
        
        # Title
        title = doc.add_heading('Hawkeye Review Feedback Summary', 1)
        
        # Metadata
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total feedback items: {len(comments_data)}')
        doc.add_paragraph('')
        
        # Group comments by section and type
        from collections import defaultdict
        section_comments = defaultdict(list)
        highlight_comments = defaultdict(list)
        
        for comment in comments_data:
            section = comment.get('section', 'Unknown Section')
            if comment.get('highlight_id') or comment.get('highlighted_text'):
                highlight_comments[section].append(comment)
            else:
                section_comments[section].append(comment)
        
        # Add each section's feedback
        for section_name in set(list(section_comments.keys()) + list(highlight_comments.keys())):
            section_heading = doc.add_heading(section_name, 2)
            
            # Regular feedback
            regular_comments = section_comments.get(section_name, [])
            if regular_comments:
                doc.add_heading('General Feedback', 3)
                for i, comment in enumerate(regular_comments, 1):
                    self._add_comment_to_summary(doc, comment, i)
            
            # Highlighted text feedback
            highlighted_comments = highlight_comments.get(section_name, [])
            if highlighted_comments:
                doc.add_heading('Text-Specific Comments', 3)
                for i, comment in enumerate(highlighted_comments, 1):
                    self._add_highlight_comment_to_summary(doc, comment, i)
            
            doc.add_paragraph('')  # Add spacing between sections
    
    def _add_comment_to_summary(self, doc, comment, index):
        """Add a regular comment to the summary"""
        # Create feedback item
        para = doc.add_paragraph(style='List Number')
        
        # Author and type
        author_run = para.add_run(f"[{comment.get('author', 'AI Feedback')}] ")
        author_run.bold = True
        
        # Risk level with color
        risk_level = comment.get('risk_level', 'Low')
        type_text = f"{comment.get('type', 'feedback').upper()} - {risk_level} Risk: "
        type_run = para.add_run(type_text)
        type_run.bold = True
        
        if risk_level == 'High':
            type_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
        elif risk_level == 'Medium':
            type_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
        else:
            type_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue
        
        # Comment text
        para.add_run(comment.get('comment', ''))
        
        doc.add_paragraph('')  # Add spacing
    
    def _add_highlight_comment_to_summary(self, doc, comment, index):
        """Add a highlighted text comment to the summary"""
        # Create feedback item
        para = doc.add_paragraph(style='List Number')
        
        # Author and type
        author_run = para.add_run(f"[{comment.get('author', 'User Highlight')}] ")
        author_run.bold = True
        
        # Highlight indicator
        highlight_run = para.add_run("📝 HIGHLIGHTED TEXT: ")
        highlight_run.bold = True
        highlight_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        
        # Risk level with color
        risk_level = comment.get('risk_level', 'Low')
        type_text = f"{comment.get('type', 'feedback').upper()} - {risk_level} Risk\n"
        type_run = para.add_run(type_text)
        type_run.bold = True
        
        if risk_level == 'High':
            type_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
        elif risk_level == 'Medium':
            type_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
        else:
            type_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue
        
        # Highlighted text
        if comment.get('highlighted_text'):
            highlighted_para = doc.add_paragraph()
            highlighted_run = highlighted_para.add_run(f'Highlighted Text: "{comment.get("highlighted_text")}"')
            highlighted_run.italic = True
            highlighted_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Comment text
        comment_para = doc.add_paragraph()
        comment_para.add_run(f'Comment: {comment.get("comment", "")}')
        
        doc.add_paragraph('')  # Add spacing

    def cleanup_temp_files(self):
        """Clean up temporary directories"""
        for temp_dir in self.temp_dirs:
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
        self.temp_dirs = []

    def __del__(self):
        """Cleanup on destruction"""
        self.cleanup_temp_files()